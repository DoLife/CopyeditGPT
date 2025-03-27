import os
import logging
from flask import Flask, render_template, request, send_file, redirect
from werkzeug.exceptions import RequestEntityTooLarge
from functions import run_editor
import docx
import requests
import webbrowser
import threading
import time

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    filename="app.log",
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    filemode="w"
)
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.config["UPLOAD_DIRECTORY"] = 'text_files/'
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024  # 16MB
app.config["ALLOWED_EXTENSIONS"] = [".txt", ".docx"]

class GlobalState:
    def __init__(self):
        self.submit_text = ""

global_state = GlobalState()

def check_ollama_status():
    """Check if Ollama server is running and the model is available"""
    try:
        # First check if server is running
        health_check = requests.get("http://localhost:11434")
        if health_check.status_code != 200:
            return False
            
        # Now check if our model exists
        response = requests.post(
            "http://localhost:11434/api/show",
            json={"name": "llama3.2-8b-instruct-128k:latest"}
        )
        return response.status_code == 200
        
    except requests.exceptions.RequestException:
        return False

def open_browser():
    """Open browser after a short delay"""
    time.sleep(1.5)  # Wait for Flask to start
    webbrowser.open('http://localhost:5000/')

@app.route('/', methods=["GET"])
def index():
    # Check Ollama status when loading the main page
    ollama_available = check_ollama_status()
    if not ollama_available:
        return render_template("error.html", 
                             error="Ollama server is not running or required model is not available. " 
                             "Please ensure Ollama is running and the model is installed.")
    return render_template("index.html")

@app.route('/upload', methods=["POST"])
def upload():
    global_state.submit_text = ""  # Reset submit text
    
    try:
        if request.form['upload'] == "upload_files":
            # Get list of uploaded files
            files = request.files.getlist('files')
            
            if not files or not files[0].filename:
                return "Must upload at least one file"

            # Process each file and combine their content
            combined_text = ""
            for file in files:
                extension = os.path.splitext(file.filename)[1].lower()
                if extension not in app.config["ALLOWED_EXTENSIONS"]:
                    return f"Cannot upload file type {extension}. Must be '.txt' or '.docx'"

                if extension == ".txt":
                    content = file.read().decode('utf-8', errors='ignore')
                    combined_text += f"\n=== File: {file.filename} ===\n\n{content}\n\n"

                elif extension == ".docx":
                    doc = docx.Document(file)
                    content = "\n".join(paragraph.text for paragraph in doc.paragraphs)
                    combined_text += f"\n=== File: {file.filename} ===\n\n{content}\n\n"

            global_state.submit_text = combined_text

        elif request.form['upload'] == "upload_text":
            text = request.form['text_box']
            if not text.strip():
                return "Text box is blank"
            global_state.submit_text = text

        return redirect('/progress')

    except RequestEntityTooLarge:
        return "File is too large. Maximum size is 16MB."
    except Exception as e:
        logger.error(f"Error in upload: {str(e)}")
        return f"An error occurred: {str(e)}"

@app.route('/progress', methods=["GET", "POST"])
def progress():
    if not global_state.submit_text:
        return redirect('/')
        
    chunk_count = (len(global_state.submit_text) // 4000) + 1
    
    if request.method == "GET": 
        return render_template(
            "progress.html",
            chunks=chunk_count,
            wait=chunk_count * 15
        )
        
    if request.method == "POST":
        try:
            run_editor(global_state.submit_text, chunk_count)
            return redirect('/results')
        except Exception as e:
            logger.error(f"Error in processing: {str(e)}")
            return render_template("error.html", error=str(e))

@app.route('/results')
def results():    
    try:
        with open("text_files/edited.txt", "r", encoding='utf-8', errors="ignore") as f:
            edited_text = f.read().split("\n")
        return render_template("results.html", text_to_display=edited_text)
    except Exception as e:
        logger.error(f"Error displaying results: {str(e)}")
        return render_template("error.html", error=str(e))

@app.route('/download')
def download(): 
    file_type = request.args.get('type')
    try:
        if file_type == "txt":
            return send_file("text_files/edited.txt", as_attachment=True)
        
        if file_type == "docx":
            edited = docx.Document()
            with open("text_files/edited.txt", "r", encoding='utf-8', errors="ignore") as f:
                edited_text = f.read().split("\n")
            
            for paragraph in edited_text:
                if paragraph.strip():  # Only add non-empty paragraphs
                    edited.add_paragraph(paragraph)
                    
            edited.save("text_files/edited.docx")
            return send_file("text_files/edited.docx", as_attachment=True)
    except Exception as e:
        logger.error(f"Error in download: {str(e)}")
        return render_template("error.html", error=str(e))

if __name__ == "__main__":
    # Start the browser opener in a separate thread
    threading.Thread(target=open_browser).start()
    # Start the Flask app
    app.run(debug=True)