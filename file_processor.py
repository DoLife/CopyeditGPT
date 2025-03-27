import os
import logging
from typing import List, Dict
import docx
from werkzeug.datastructures import FileStorage

logger = logging.getLogger(__name__)

class FileProcessor:
    def __init__(self, upload_dir: str):
        self.upload_dir = upload_dir
        self.processed_files: Dict[str, str] = {}  # filename -> edited_filename mapping
        self.current_file: str = ""
        self.total_files: int = 0
        self.processed_count: int = 0
        
    def process_files(self, files: List[FileStorage]) -> List[Dict[str, str]]:
        """Process multiple files and return their contents"""
        self.total_files = len(files)
        self.processed_count = 0
        
        file_contents = []
        for file in files:
            if not file or not file.filename:
                continue
                
            try:
                self.current_file = file.filename
                content = self._process_single_file(file)
                if content:
                    file_contents.append({
                        'filename': file.filename,
                        'content': content,
                        'output_filename': self._get_output_filename(file.filename)
                    })
                self.processed_count += 1
            except Exception as e:
                logger.error(f"Error processing file {file.filename}: {str(e)}")
                continue
                
        return file_contents
    
    def _process_single_file(self, file: FileStorage) -> str:
        """Process a single file and return its content"""
        extension = os.path.splitext(file.filename)[1].lower()
        
        if extension == '.txt':
            return file.read().decode('utf-8', errors='ignore')
            
        elif extension == '.docx':
            doc = docx.Document(file)
            return '\n'.join(paragraph.text for paragraph in doc.paragraphs)
            
        return ""
    
    def _get_output_filename(self, input_filename: str) -> str:
        """Generate output filename for edited content"""
        base, ext = os.path.splitext(input_filename)
        output_filename = f"{base}_edited{ext}"
        self.processed_files[input_filename] = output_filename
        return output_filename
    
    def save_edited_content(self, filename: str, content: str) -> str:
        """Save edited content to a file"""
        output_filename = self.processed_files.get(filename)
        if not output_filename:
            return ""
            
        output_path = os.path.join(self.upload_dir, output_filename)
        extension = os.path.splitext(filename)[1].lower()
        
        try:
            if extension == '.txt':
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                    
            elif extension == '.docx':
                doc = docx.Document()
                for paragraph in content.split('\n'):
                    if paragraph.strip():
                        doc.add_paragraph(paragraph)
                doc.save(output_path)
                
            return output_filename
            
        except Exception as e:
            logger.error(f"Error saving edited content for {filename}: {str(e)}")
            return ""
    
    def get_progress(self) -> Dict[str, any]:
        """Get current processing progress"""
        return {
            'current_file': self.current_file,
            'processed_count': self.processed_count,
            'total_files': self.total_files,
            'percentage': (self.processed_count / self.total_files * 100) if self.total_files > 0 else 0
        }